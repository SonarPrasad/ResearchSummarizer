import os
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
import google.generativeai as genai
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Utility functions

def get_pdf_text(pdf_docs):
    """Extract text from uploaded PDF files."""
    text = ""
    for pdf in pdf_docs:
        pdf_read = PdfReader(pdf)
        for page in pdf_read.pages:
            text += page.extract_text()
    return text

def get_text_chunks(text):
    """Split text into smaller chunks."""
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=5000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks

def get_vector_store(pdf_docs):
    """Create and save a vector store from text chunks of multiple PDF files."""
    all_text_chunks = []
    
    for pdf in pdf_docs:
        pdf_path = pdf.file.path  # Get the file path
        raw_text = get_pdf_text([pdf_path])  # Extract text
        text_chunks = get_text_chunks(raw_text)  # Split into chunks
        all_text_chunks.extend(text_chunks)  # Add to the combined list

    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(all_text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")  # Save to local file

def get_conversational_chain():
    """Get a QA chain for answering questions using generative AI."""

    prompt_template = """
    Answer the question as detailed as possible from the provided context. If the answer is not in the
    provided context, just say, "Answer is not available in the context." Please avoid guessing.\n\n
    Context:\n {context}?\n
    Question:\n {question}\n
    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

    return chain

def process_user_input(user_question):
    """Process user question and return the formatted response."""
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search(user_question, k=5)

    chain = get_conversational_chain()

    response = chain.invoke({
        "input_documents": docs,
        "question": user_question
    })

    output_text = response.get("output_text", "")
    formatted_text = output_text.replace("\n", "<br>")
    return formatted_text

def generate_word_content(title, user_input):
    """
    Generate a Word file based on user input and AI-generated content from uploaded PDFs.

    Args:
    - title (str): The title of the Word file.
    - user_input (str): The input text from the user for generating content, separated by commas.

    Returns:
    - file_path (str): The file path of the generated Word document.
    """
    # Create a new Word document
    document = Document()

    # Add a title
    document.add_heading(title, level=1)

    # Get the vector store and conversational chain
    vector_store_path = "faiss_index"
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.load_local(vector_store_path, embeddings, allow_dangerous_deserialization=True)
    chain = get_conversational_chain()

    # Split the user input by commas to handle multiple topics
    topics = [topic.strip() for topic in user_input.split(',')]

    for topic in topics:
        # Search the vector store with the current topic
        docs = vector_store.similarity_search(topic, k=5)

        # Log the found documents for debugging
        print(f"Documents found for topic '{topic}': {[doc.page_content[:500] for doc in docs]}")  # Print first 500 characters of each document

        # Check if any documents were found
        if not docs:
            ai_content = "No relevant content found for your question."
        else:
            # Generate a response using the chain with context from the found documents
            content = "\n\n".join([doc.page_content for doc in docs])  # Join document contents

            response = chain.invoke({
                "input_documents": docs,
                "question": topic
            })

            print(f"Response for topic '{topic}': {response}")

            ai_content = response.get("output_text", "Answer is not available in the context.")

        # Add topic heading and AI-generated content to the document
        document.add_heading(topic, level=2)
        document.add_paragraph("User Input:")
        document.add_paragraph(topic)
        document.add_paragraph("\nGenerated Content:")
        document.add_paragraph(ai_content)

    # Save the document to a file
    file_name = f"{title.replace(' ', '_')}.docx"
    file_path = os.path.join('generated_files', file_name)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)  # Create directory if not exists
    document.save(file_path)

    return file_path